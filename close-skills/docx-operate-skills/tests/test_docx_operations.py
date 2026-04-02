"""
Unit Tests for DOCX Operate Skills

This module contains unit tests for the DOCX manipulation
functions provided in docx_operations.py.
"""
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.text.paragraph import Paragraph

from docx_operations import (
    is_heading_rank,
    normalize_heading_text,
    replace_section_by_title,
    clear_section_content,
    get_subtitles_under_heading,
    add_hyperlink_to_heading,
    insert_heading_after_subtitles,
    _parse_heading_level,
    _iter_blocks,
    _find_section,
)


class TestHeadingLevelDetection(unittest.TestCase):
    """Tests for heading level detection functions."""

    def test_parse_heading_level_english(self):
        """Test parsing English heading style names."""
        self.assertEqual(_parse_heading_level("Heading 1"), 1)
        self.assertEqual(_parse_heading_level("Heading 2"), 2)
        self.assertEqual(_parse_heading_level("Heading 9"), 9)

    def test_parse_heading_level_chinese(self):
        """Test parsing Chinese heading style names."""
        self.assertEqual(_parse_heading_level("标题 1"), 1)
        self.assertEqual(_parse_heading_level("标题 2"), 2)
        self.assertEqual(_parse_heading_level("标题1"), 1)

    def test_parse_heading_level_subtitle(self):
        """Test parsing Subtitle style names."""
        self.assertEqual(_parse_heading_level("Subtitle"), 1)
        self.assertEqual(_parse_heading_level("副标题"), 1)

    def test_parse_heading_level_invalid(self):
        """Test parsing invalid style names."""
        self.assertIsNone(_parse_heading_level("Normal"))
        self.assertIsNone(_parse_heading_level("Body Text"))
        self.assertIsNone(_parse_heading_level(""))
        self.assertIsNone(_parse_heading_level(None))

    def test_is_heading_rank_with_document(self):
        """Test is_heading_rank with actual document paragraphs."""
        doc = Document()
        
        heading1 = doc.add_heading("Test Heading 1", level=1)
        self.assertEqual(is_heading_rank(heading1), 1)
        
        heading2 = doc.add_heading("Test Heading 2", level=2)
        self.assertEqual(is_heading_rank(heading2), 2)
        
        normal = doc.add_paragraph("Normal paragraph")
        self.assertIsNone(is_heading_rank(normal))


class TestHeadingTextNormalization(unittest.TestCase):
    """Tests for heading text normalization."""

    def test_numeric_prefix_removal(self):
        """Test removal of numeric prefixes."""
        self.assertEqual(normalize_heading_text("1. Introduction"), "Introduction")
        self.assertEqual(normalize_heading_text("1.2 Methods"), "Methods")
        self.assertEqual(normalize_heading_text("1.2.3 Results"), "Results")
        self.assertEqual(normalize_heading_text("2．Discussion"), "Discussion")

    def test_chinese_numeric_prefix_removal(self):
        """Test removal of Chinese numeric prefixes."""
        self.assertEqual(normalize_heading_text("一、引言"), "引言")
        self.assertEqual(normalize_heading_text("二、方法"), "方法")
        self.assertEqual(normalize_heading_text("三、结果"), "结果")

    def test_parenthesized_prefix_removal(self):
        """Test removal of parenthesized prefixes."""
        self.assertEqual(normalize_heading_text("(1) Introduction"), "Introduction")
        self.assertEqual(normalize_heading_text("（1）引言"), "引言")
        self.assertEqual(normalize_heading_text("(1.1) Methods"), "Methods")

    def test_empty_and_whitespace(self):
        """Test handling of empty and whitespace strings."""
        self.assertEqual(normalize_heading_text(""), "")
        self.assertEqual(normalize_heading_text("   "), "")
        self.assertEqual(normalize_heading_text("  Introduction  "), "Introduction")


class TestIterBlocks(unittest.TestCase):
    """Tests for document block iteration."""

    def test_iter_blocks_paragraphs(self):
        """Test iteration over paragraphs."""
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        doc.add_paragraph("Paragraph 3")

        blocks = list(_iter_blocks(doc))
        self.assertEqual(len(blocks), 3)
        
        for element, block in blocks:
            self.assertIsInstance(block, Paragraph)

    def test_iter_blocks_with_tables(self):
        """Test iteration over paragraphs and tables."""
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_table(rows=2, cols=2)
        doc.add_paragraph("Paragraph 2")

        blocks = list(_iter_blocks(doc))
        self.assertEqual(len(blocks), 3)


class TestFindSection(unittest.TestCase):
    """Tests for section finding."""

    def setUp(self):
        """Create a test document."""
        self.doc = Document()
        self.doc.add_heading("Introduction", level=1)
        self.doc.add_paragraph("Intro content 1")
        self.doc.add_paragraph("Intro content 2")
        self.doc.add_heading("Methods", level=1)
        self.doc.add_paragraph("Methods content 1")
        self.doc.add_heading("Results", level=1)
        self.doc.add_paragraph("Results content 1")

    def test_find_section_found(self):
        """Test finding existing section."""
        heading, nodes = _find_section(self.doc, "Introduction")
        self.assertIsNotNone(heading)
        self.assertEqual(len(nodes), 2)

    def test_find_section_not_found(self):
        """Test finding non-existent section."""
        heading, nodes = _find_section(self.doc, "NonExistent")
        self.assertIsNone(heading)
        self.assertEqual(len(nodes), 0)

    def test_find_section_normalized(self):
        """Test finding section with normalized heading."""
        doc = Document()
        doc.add_heading("1. Introduction", level=1)
        doc.add_paragraph("Content")

        heading, nodes = _find_section(doc, "Introduction")
        self.assertIsNotNone(heading)
        self.assertEqual(len(nodes), 1)

    def test_find_section_last_match(self):
        """Test finding last matched section."""
        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_paragraph("Content A")
        doc.add_heading("Another", level=1)
        doc.add_paragraph("Content B")
        doc.add_heading("Section", level=1)
        doc.add_paragraph("Content C")

        heading, nodes = _find_section(doc, "Section", use_last=True)
        self.assertIsNotNone(heading)
        self.assertEqual(len(nodes), 1)


class TestGetSubtitlesUnderHeading(unittest.TestCase):
    """Tests for getting subtitles under a heading."""

    def setUp(self):
        """Create test document file."""
        self.temp_dir = tempfile.mkdtemp()
        self.doc_path = os.path.join(self.temp_dir, "test.docx")
        
        doc = Document()
        doc.add_heading("Main Section", level=1)
        doc.add_heading("Subsection 1", level=2)
        doc.add_paragraph("Content 1")
        doc.add_heading("Subsection 2", level=2)
        doc.add_paragraph("Content 2")
        doc.add_heading("Subsection 3", level=2)
        doc.add_paragraph("Content 3")
        doc.add_heading("Another Main", level=1)
        doc.save(self.doc_path)

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_get_subtitles(self):
        """Test getting subtitles under a heading."""
        subtitles = get_subtitles_under_heading(self.doc_path, "Main Section")
        self.assertEqual(len(subtitles), 3)
        
        original_texts = [orig for orig, _ in subtitles]
        self.assertIn("Subsection 1", original_texts)
        self.assertIn("Subsection 2", original_texts)
        self.assertIn("Subsection 3", original_texts)

    def test_get_subtitles_not_found(self):
        """Test getting subtitles for non-existent heading."""
        with self.assertRaises(ValueError):
            get_subtitles_under_heading(self.doc_path, "NonExistent")


class TestClearSectionContent(unittest.TestCase):
    """Tests for clearing section content."""

    def setUp(self):
        """Create test document file."""
        self.temp_dir = tempfile.mkdtemp()
        self.doc_path = os.path.join(self.temp_dir, "test.docx")

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_clear_section_content(self):
        """Test clearing content under a heading."""
        doc = Document()
        doc.add_heading("Test Section", level=1)
        doc.add_paragraph("Content to clear 1")
        doc.add_paragraph("Content to clear 2")
        doc.add_heading("Next Section", level=1)
        doc.save(self.doc_path)

        clear_section_content(self.doc_path, "Test Section")

        doc = Document(self.doc_path)
        paragraphs = [p.text for p in doc.paragraphs]
        self.assertNotIn("Content to clear 1", paragraphs)
        self.assertNotIn("Content to clear 2", paragraphs)

    def test_clear_with_exclusions(self):
        """Test clearing content with exclusions."""
        doc = Document()
        doc.add_heading("Main", level=1)
        doc.add_paragraph("Content 1")
        doc.add_heading("Keep This", level=2)
        doc.add_paragraph("Important content")
        doc.add_heading("Remove This", level=2)
        doc.add_paragraph("Unimportant content")
        doc.save(self.doc_path)

        clear_section_content(
            self.doc_path, 
            "Main", 
            keep_subtitles=["Keep This"]
        )

        doc = Document(self.doc_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("Important content", texts)
        self.assertNotIn("Unimportant content", texts)


class TestReplaceSectionByTitle(unittest.TestCase):
    """Tests for section replacement."""

    def setUp(self):
        """Create test document files."""
        self.temp_dir = tempfile.mkdtemp()
        self.source_path = os.path.join(self.temp_dir, "source.docx")
        self.target_path = os.path.join(self.temp_dir, "target.docx")

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_replace_section_basic(self):
        """Test basic section replacement."""
        source_doc = Document()
        source_doc.add_heading("Section", level=1)
        source_doc.add_paragraph("New content from source")
        source_doc.save(self.source_path)

        target_doc = Document()
        target_doc.add_heading("Section", level=1)
        target_doc.add_paragraph("Old content to replace")
        target_doc.add_heading("Next", level=1)
        target_doc.save(self.target_path)

        replace_section_by_title(self.source_path, self.target_path, "Section")

        result_doc = Document(self.target_path)
        texts = [p.text for p in result_doc.paragraphs]
        self.assertIn("New content from source", texts)
        self.assertNotIn("Old content to replace", texts)

    def test_replace_with_title_mapping(self):
        """Test section replacement with title mapping."""
        source_doc = Document()
        source_doc.add_heading("Source Title", level=1)
        source_doc.add_paragraph("Content from source")
        source_doc.save(self.source_path)

        target_doc = Document()
        target_doc.add_heading("Target Title", level=1)
        target_doc.add_paragraph("Old content")
        target_doc.save(self.target_path)

        replace_section_by_title(
            self.source_path,
            self.target_path,
            "Source Title",
            title_mapping={"Source Title": "Target Title"}
        )

        result_doc = Document(self.target_path)
        texts = [p.text for p in result_doc.paragraphs]
        self.assertIn("Content from source", texts)

    def test_replace_last_matched(self):
        """Test replacement uses the last matched heading."""
        source_doc = Document()
        source_doc.add_heading("Section", level=1)
        source_doc.add_paragraph("New content")
        source_doc.save(self.source_path)

        target_doc = Document()
        target_doc.add_heading("Section", level=1)
        target_doc.add_paragraph("First section content")
        target_doc.add_heading("Other", level=1)
        target_doc.add_paragraph("Other content")
        target_doc.add_heading("Section", level=1)
        target_doc.add_paragraph("Second section content")
        target_doc.save(self.target_path)

        replace_section_by_title(
            self.source_path,
            self.target_path,
            "Section",
            use_last_match=True
        )

        result_doc = Document(self.target_path)
        paragraphs = [p.text for p in result_doc.paragraphs]
        
        self.assertIn("First section content", paragraphs)
        self.assertIn("New content", paragraphs)
        self.assertNotIn("Second section content", paragraphs)


class TestAddHyperlinkToHeading(unittest.TestCase):
    """Tests for adding hyperlinks to headings."""

    def setUp(self):
        """Create test document file."""
        self.temp_dir = tempfile.mkdtemp()
        self.doc_path = os.path.join(self.temp_dir, "test.docx")

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_add_hyperlink_basic(self):
        """Test adding a basic hyperlink."""
        doc = Document()
        doc.add_heading("Target Heading", level=1)
        doc.add_paragraph("This is a reference to Target Heading.")
        doc.save(self.doc_path)

        count = add_hyperlink_to_heading(
            self.doc_path,
            "Target Heading",
            "Target Heading"
        )

        self.assertGreater(count, 0)

    def test_add_hyperlink_heading_not_found(self):
        """Test error when heading not found."""
        doc = Document()
        doc.add_paragraph("Some content")
        doc.save(self.doc_path)

        with self.assertRaises(ValueError):
            add_hyperlink_to_heading(
                self.doc_path,
                "phrase",
                "NonExistent Heading"
            )


class TestInsertHeadingAfterSubtitles(unittest.TestCase):
    """Tests for inserting headings."""

    def setUp(self):
        """Create test document file."""
        self.temp_dir = tempfile.mkdtemp()
        self.doc_path = os.path.join(self.temp_dir, "test.docx")

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_insert_heading(self):
        """Test inserting a heading after subtitles."""
        doc = Document()
        doc.add_heading("Main Section", level=1)
        doc.add_heading("Subsection A", level=2)
        doc.add_paragraph("Content A")
        doc.add_heading("Subsection B", level=2)
        doc.add_paragraph("Content B")
        doc.save(self.doc_path)

        insert_heading_after_subtitles(
            self.doc_path,
            "New Subsection",
            "Main Section",
            level=2
        )

        result_doc = Document(self.doc_path)
        headings = [p.text for p in result_doc.paragraphs if is_heading_rank(p)]
        self.assertIn("New Subsection", headings)


class TestInputValidation(unittest.TestCase):
    """Tests for input validation across functions."""

    def test_replace_section_empty_title(self):
        """Test that empty title raises ValueError."""
        with self.assertRaises(ValueError):
            replace_section_by_title("source.docx", "target.docx", "")

    def test_clear_section_empty_path(self):
        """Test that empty file path raises ValueError."""
        with self.assertRaises(ValueError):
            clear_section_content("", "Title")

    def test_get_subtitles_empty_title(self):
        """Test that empty parent title raises ValueError."""
        with self.assertRaises(ValueError):
            get_subtitles_under_heading("file.docx", "")

    def test_add_hyperlink_empty_phrase(self):
        """Test that empty phrase raises ValueError."""
        with self.assertRaises(ValueError):
            add_hyperlink_to_heading("file.docx", "", "Title")


if __name__ == "__main__":
    unittest.main(verbosity=2)

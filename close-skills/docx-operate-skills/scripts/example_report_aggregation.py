"""
Example: Document Structure Analysis

This script demonstrates how to analyze document structure using
subtitle detection and heading operations.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx_operations import (
    get_subtitles_under_heading,
    is_heading_rank,
    normalize_heading_text,
)
from docx import Document


def example_list_subtitles():
    """
    List all subtitles under a parent heading.
    """
    file_path = "template.docx"
    parent_title = "Results"

    try:
        subtitles = get_subtitles_under_heading(
            file_path=file_path,
            parent_title=parent_title,
        )
        print(f"Subtitles under '{parent_title}':")
        for original, normalized in subtitles:
            print(f"  - {original} (normalized: {normalized})")
    except ValueError as e:
        print(f"Error: {e}")


def example_analyze_document_structure():
    """
    Analyze the complete structure of a document.
    """
    file_path = "report.docx"

    try:
        doc = Document(file_path)
        print("Document Structure Analysis:")
        print("-" * 40)
        
        for para in doc.paragraphs:
            level = is_heading_rank(para)
            if level is not None:
                indent = "  " * (level - 1)
                text = para.text.strip()
                normalized = normalize_heading_text(text)
                print(f"{indent}[H{level}] {text}")
                if normalized != text:
                    print(f"{indent}      -> {normalized}")
                    
    except Exception as e:
        print(f"Error: {e}")


def example_find_matching_sections():
    """
    Find sections that match a specific pattern.
    """
    file_path = "report.docx"
    search_term = "analysis"

    try:
        doc = Document(file_path)
        matches = []
        
        for para in doc.paragraphs:
            level = is_heading_rank(para)
            if level is not None:
                normalized = normalize_heading_text(para.text).lower()
                if search_term.lower() in normalized:
                    matches.append((level, para.text.strip()))
        
        print(f"Sections containing '{search_term}':")
        for level, text in matches:
            print(f"  [H{level}] {text}")
            
    except Exception as e:
        print(f"Error: {e}")


def example_verify_template_structure():
    """
    Verify template structure before processing.
    """
    template_path = "template.docx"
    required_sections = ["Introduction", "Methods", "Results", "Conclusion"]

    print("Verifying template structure...")
    all_found = True
    
    for section in required_sections:
        try:
            subtitles = get_subtitles_under_heading(
                file_path=template_path,
                parent_title=section,
            )
            print(f"  {section}: {len(subtitles)} subtitle(s)")
        except ValueError:
            print(f"  {section}: NOT FOUND")
            all_found = False
    
    if all_found:
        print("\n[OK] All required sections found")
    else:
        print("\n[WARNING] Some sections are missing")


if __name__ == "__main__":
    print("=" * 60)
    print("Document Structure Analysis Examples")
    print("=" * 60)

    print("\n1. List Subtitles:")
    print("   See example_list_subtitles()")

    print("\n2. Analyze Document Structure:")
    print("   See example_analyze_document_structure()")

    print("\n3. Find Matching Sections:")
    print("   See example_find_matching_sections()")

    print("\n4. Verify Template Structure:")
    print("   See example_verify_template_structure()")

    print("\n" + "=" * 60)
    print("Note: Ensure document files exist before running examples.")
    print("=" * 60)

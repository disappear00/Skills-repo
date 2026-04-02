"""
DOCX Operations Module - Simplified

A lightweight module for manipulating Microsoft Word (DOCX) documents.
Provides core utilities for section replacement, content management, and hyperlinks.
"""
import copy
import hashlib
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table


__all__ = [
    "replace_section_by_title",
    "clear_section_content",
    "get_subtitles_under_heading",
    "add_hyperlink_to_heading",
    "insert_heading_after_subtitles",
    "is_heading_rank",
    "normalize_heading_text",
]


def _parse_heading_level(style_name: Optional[str]) -> Optional[int]:
    """Parse heading level from style name (supports English and Chinese)."""
    if not style_name:
        return None
    if style_name.startswith(("Subtitle", "副标题")):
        return 1
    for prefix in ("Heading", "标题"):
        if style_name.startswith(prefix):
            match = re.search(rf"{prefix}\s*(\d+)", style_name)
            if match:
                return int(match.group(1))
    return None


def is_heading_rank(paragraph: Paragraph) -> Optional[int]:
    """Get heading level from paragraph's style name (1-9 or None)."""
    try:
        style = paragraph.style
        if style and style.name:
            return _parse_heading_level(style.name)
    except (KeyError, AttributeError):
        pass
    return None


_HEADING_PATTERNS = [
    re.compile(r"^\s*\d+(?:\.\d+)*[．\.\-、,：:\s]*"),
    re.compile(r"^\s*[（(]\d+(?:\.\d+)*[）)][．\.\-、,：:\s]*"),
    re.compile(r"^\s*[一二三四五六七八九十百千零]+[、.．：:\s-]*"),
    re.compile(r"^\s*[（(][一二三四五六七八九十百千零]+[)）][、.．：:\s-]*"),
]


def normalize_heading_text(text: str) -> str:
    """Normalize heading text by stripping leading numbering."""
    if not text:
        return ""
    result = text.strip()
    for pattern in _HEADING_PATTERNS:
        result = pattern.sub("", result, count=1)
    return result.strip()


def _iter_blocks(doc: Document):
    """Iterate over document body blocks, yielding (element, block)."""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield child, Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield child, Table(child, doc)


def _find_section(doc: Document, heading_text: str, use_last: bool = False) -> Tuple[Optional[CT_P], List]:
    """Find heading and its content nodes."""
    heading_key = normalize_heading_text(heading_text) or heading_text.strip()
    matches = []
    
    for element, block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            if normalize_heading_text(block.text.strip()) == heading_key:
                matches.append((element, is_heading_rank(block)))
    
    if not matches:
        return None, []
    
    target_element, target_level = matches[-1] if use_last else matches[0]
    target_level = target_level or 1
    
    nodes = []
    capture = False
    for element, block in _iter_blocks(doc):
        if element is target_element:
            capture = True
            continue
        if capture:
            if isinstance(block, Paragraph):
                level = is_heading_rank(block)
                if level is not None and level <= target_level:
                    break
            nodes.append(element)
    
    return target_element, nodes


def replace_section_by_title(
    source_file: str,
    target_file: str,
    title: str,
    title_mapping: Optional[Dict[str, str]] = None,
    use_last_match: bool = False,
) -> None:
    """
    Replace content under a heading in target file with content from source file.
    
    Args:
        source_file: Path to source document.
        target_file: Path to target document (modified in place).
        title: Heading text to find and replace.
        title_mapping: Optional mapping {source_title: target_title}.
        use_last_match: If True, use last matching heading when duplicates exist.
    """
    if not title or not title.strip():
        raise ValueError("title must be a non-empty string")

    source_title = title.strip()
    target_title = (title_mapping or {}).get(source_title, source_title)

    source_doc = Document(source_file)
    target_doc = Document(target_file)

    src_heading, src_nodes = _find_section(source_doc, source_title)
    if src_heading is None:
        raise ValueError(f"Heading '{source_title}' not found in source document.")

    tgt_heading, tgt_nodes = _find_section(target_doc, target_title, use_last_match)
    if tgt_heading is None:
        raise ValueError(f"Heading '{target_title}' not found in target document.")

    parent = tgt_heading.getparent()
    if parent is None:
        raise ValueError("Cannot locate target heading parent.")

    for node in tgt_nodes:
        parent.remove(node)

    insert_after = tgt_heading
    for node in src_nodes:
        cloned = copy.deepcopy(node)
        insert_after.addnext(cloned)
        insert_after = cloned

    target_doc.save(target_file)


def clear_section_content(
    file_path: str,
    title: str,
    keep_subtitles: Optional[List[str]] = None,
) -> None:
    """
    Remove content under a heading while keeping the heading itself.
    
    Args:
        file_path: Path to document (modified in place).
        title: Heading text whose content should be removed.
        keep_subtitles: List of subheading names whose content to preserve.
    """
    if not title or not title.strip():
        raise ValueError("title must be a non-empty string")

    doc = Document(file_path)
    heading, nodes = _find_section(doc, title.strip())
    
    if heading is None:
        raise ValueError(f"Heading '{title}' not found in document.")

    if not nodes:
        doc.save(file_path)
        return

    heading_level = is_heading_rank(Paragraph(heading, doc)) or 1
    keep_set = {normalize_heading_text(t) for t in (keep_subtitles or [])}
    
    nodes_to_remove = []
    skip_until_level = None
    
    for node in nodes:
        if isinstance(node, CT_P):
            para = Paragraph(node, doc)
            level = is_heading_rank(para)
            
            if level is not None:
                if skip_until_level is not None and level <= skip_until_level:
                    skip_until_level = None
                norm_name = normalize_heading_text(para.text.strip())
                if norm_name in keep_set:
                    skip_until_level = level
                    continue
            
            if skip_until_level is not None and (level is None or level > skip_until_level):
                continue
        
        if skip_until_level is None:
            nodes_to_remove.append(node)

    for node in nodes_to_remove:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)

    doc.save(file_path)


def get_subtitles_under_heading(
    file_path: str,
    parent_title: str,
) -> List[Tuple[str, str]]:
    """
    Get all direct child headings under a specified parent heading.
    
    Args:
        file_path: Path to document.
        parent_title: Parent heading text.
    
    Returns:
        List of (original_text, normalized_text) tuples.
    """
    if not parent_title or not parent_title.strip():
        raise ValueError("parent_title must be a non-empty string")

    doc = Document(file_path)
    heading, nodes = _find_section(doc, parent_title.strip())
    
    if heading is None:
        raise ValueError(f"Heading '{parent_title}' not found in document.")

    parent_level = is_heading_rank(Paragraph(heading, doc)) or 1
    child_level = parent_level + 1

    subtitles: List[Tuple[str, str]] = []
    for node in nodes:
        if isinstance(node, CT_P):
            para = Paragraph(node, doc)
            level = is_heading_rank(para)
            if level == child_level:
                original = para.text.strip()
                normalized = normalize_heading_text(original)
                if normalized:
                    subtitles.append((original, normalized))

    return subtitles


def _make_bookmark_name(title: str) -> str:
    """Generate a unique bookmark name from title."""
    hash_val = hashlib.md5(title.encode('utf-8')).hexdigest()[:8]
    return f"_bm_{hash_val}"


def _add_bookmark(element: CT_P, name: str) -> None:
    """Add bookmark to paragraph element."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), name)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    
    element.insert(0, start)
    element.append(end)


def _add_hyperlink(paragraph: Paragraph, phrase: str, bookmark: str) -> int:
    """Add hyperlinks for phrase occurrences in paragraph."""
    count = 0
    phrase_lower = phrase.lower()
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        pos = text.lower().find(phrase_lower)
        if pos == -1:
            continue
        
        before = text[:pos]
        after = text[pos + len(phrase):]
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark)
        
        new_run = OxmlElement('w:r')
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None:
            new_run.append(copy.deepcopy(rPr))
        
        t = OxmlElement('w:t')
        t.text = phrase
        new_run.append(t)
        hyperlink.append(new_run)
        
        run.text = before
        run_idx = list(paragraph._p).index(run._r)
        paragraph._p.insert(run_idx + 1, hyperlink)
        
        if after:
            after_run = OxmlElement('w:r')
            if rPr is not None:
                after_run.append(copy.deepcopy(rPr))
            after_t = OxmlElement('w:t')
            after_t.text = after
            after_run.append(after_t)
            paragraph._p.insert(run_idx + 2, after_run)
        
        count += 1

    return count


def add_hyperlink_to_heading(
    file_path: str,
    phrase: str,
    target_heading: str,
) -> int:
    """
    Add internal hyperlinks from text to a target heading.
    
    Args:
        file_path: Path to document (modified in place).
        phrase: Text to convert to hyperlinks.
        target_heading: Heading text to link to.
    
    Returns:
        Number of hyperlinks added.
    """
    if not phrase or not phrase.strip():
        raise ValueError("phrase must be a non-empty string")
    if not target_heading or not target_heading.strip():
        raise ValueError("target_heading must be a non-empty string")

    doc = Document(file_path)
    heading, _ = _find_section(doc, target_heading.strip())
    
    if heading is None:
        raise ValueError(f"Heading '{target_heading}' not found in document.")

    bookmark = _make_bookmark_name(target_heading.strip())
    _add_bookmark(heading, bookmark)

    count = 0
    for element, block in _iter_blocks(doc):
        if isinstance(block, Paragraph) and element is not heading:
            count += _add_hyperlink(block, phrase.strip(), bookmark)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _add_hyperlink(para, phrase.strip(), bookmark)

    doc.save(file_path)
    return count


def insert_heading_after_subtitles(
    file_path: str,
    heading_text: str,
    target_title: str,
    level: int = 2,
) -> None:
    """
    Insert a new heading after all subheadings of a target heading.
    
    Args:
        file_path: Path to document (modified in place).
        heading_text: Text for the new heading.
        target_title: Parent heading to insert under.
        level: Heading level for the new heading (default: 2).
    """
    if not heading_text or not heading_text.strip():
        raise ValueError("heading_text must be a non-empty string")
    if not target_title or not target_title.strip():
        raise ValueError("target_title must be a non-empty string")

    doc = Document(file_path)
    heading, nodes = _find_section(doc, target_title.strip())
    
    if heading is None:
        raise ValueError(f"Heading '{target_title}' not found in document.")

    parent_level = is_heading_rank(Paragraph(heading, doc)) or 1
    
    insert_after = heading
    for node in nodes:
        if isinstance(node, CT_P):
            para = Paragraph(node, doc)
            node_level = is_heading_rank(para)
            if node_level is not None and node_level <= parent_level:
                break
        insert_after = node

    new_heading = doc.add_heading(heading_text.strip(), level=level)
    new_heading._p.getparent().remove(new_heading._p)
    insert_after.addnext(new_heading._p)

    doc.save(file_path)
